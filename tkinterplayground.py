import tkinter
import tkinter as tk
from tkinter import *
import docx
from tkinter import filedialog
from docx import Document


def openFile():
    filepath = filedialog.askopenfilename(parent=root)  # Returns a string which is the filepath where the file is located
    return filepath


def word_document_input(file):
    doc = docx.Document(file)
    completedText = []
    for paragraph in doc.paragraphs:
        completedText.append(paragraph.text)
    new_string = ''.join(map(str, completedText))
    splitted = new_string.split(' ')
    first_name = 'First'
    last_name = 'Last'
    clone_name ='Clone'
    new_test_string =[]
    print(splitted)
    return splitted


# browser_btn = tk.Button(root, textvariable=browser_text, command=lambda:word_document_input(openFile()), font="Raleway", bg="#20bebe", fg="white", height=2, width=15)
class Main_window:
    def openFile(self):
        filepath = filedialog.askopenfilename(
            parent=root)  # Returns a string which is the filepath where the file is located
        return filepath

    def word_document_input(self,file):
        doc = docx.Document(file)
        completedText = []
        for paragraph in doc.paragraphs:
            completedText.append(paragraph.text)
        new_string = ''.join(map(str, completedText))
        splitted = new_string.split(':')
        print(splitted)
        return splitted
    def __init__(self, master):
        self.master = master
        self.master.geometry("400x400")
        self.frame = tk.Frame(self.master)
        self.new_button("New Starter", "2", NewStarter)
        self.new_button("Leaver", "3", Leaver)
        self.new_button("Mover", "4", Mover)
        self.frame.pack()

    def new_button(self, text, number, _class):
        tkinter.Button(self.frame, text=text, command=lambda: self.new_window(number, _class), font="Raleway",
                       bg="#20bebe", fg="white", height=2, width=15).pack()

    def new_window(self, number, _class):
        self.new = tkinter.Toplevel(self.master)
        _class(self.new, number)


class NewStarter:
    def __init__(self, master, number):
        self.master = master
        self.master.geometry("400x400")
        self.frame = tk.Frame(self.master)
        self.browse = tk.Button(self.frame, text="Browse",
                                command=self.browse_window,font="Raleway",bg="Black",fg="white",height=1,width=10)
        self.browse.pack()
        self.frame.pack()

    def browse_window(self, ):
        automatic_input = word_document_input(openFile())
        return automatic_input


class Leaver:
    def __init__(self, master, number):
        self.master = master
        self.master.geometry("400x400")
        self.frame = tk.Frame(self.master)
        self.browse = tk.Button(self.frame, text="Browse",
                                command=self.browse_window,font="Raleway",bg="Black",fg="white",height=1,width=10)
        self.browse.pack()
        self.frame.pack()

    def browse_window(self):
        return word_document_input(openFile())


class Mover:
    def __init__(self, master, number):
        self.master = master
        self.master.geometry("400x400")
        self.frame = tk.Frame(self.master)
        self.browse = tk.Button(self.frame, text="Browse",
                                command=self.browse_window,font="Raleway",bg="Black",fg="white",height=1,width=10)
        self.browse.pack()
        self.frame.pack()

    def browse_window(self):
        return word_document_input(openFile())


root = tk.Tk()
app = Main_window(root)
root.mainloop()